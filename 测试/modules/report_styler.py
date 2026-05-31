"""
交付报告美化模块
生成内部审计底稿（Word）与客户交付报告（HTML）
融入安永品牌规范与文档化要求
"""

import os
import sys
import io
from datetime import datetime
from typing import Dict, Any, List, Optional
import pandas as pd

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import EY_DISCLAIMER, EY_REVIEW_NOTICE, HAS_DOCX


def generate_word_report(
    df: pd.DataFrame,
    anomalies: pd.DataFrame,
    trend: Dict[str, Any],
    insight: Dict[str, Any],
    evidence_chain: List[Dict[str, Any]],
    project_info: Dict[str, str],
) -> io.BytesIO:
    """生成 Word 格式的审计底稿报告"""
    if not HAS_DOCX:
        raise ImportError("python-docx 未安装，请运行：pip install python-docx")

    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 封面
    title = doc.add_heading("审计数据洞察报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x00, 0x30, 0x87)

    doc.add_paragraph()
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Light Grid Accent 1"
    meta_data = [
        ("项目名称", project_info.get("name", "未命名")),
        ("审计期间", project_info.get("period", "未指定")),
        ("生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("生成人", project_info.get("preparer", "智能审计系统")),
    ]
    for i, (k, v) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = k
        meta_table.rows[i].cells[1].text = v

    doc.add_paragraph()

    # 职业怀疑提示
    doc.add_heading("职业怀疑与独立性声明", level=1)
    p = doc.add_paragraph(EY_REVIEW_NOTICE)
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)

    # 数据概览
    doc.add_heading("1. 数据概览", level=1)
    doc.add_paragraph(f"原始数据总行数：{len(df):,}")
    numeric_cols = df.select_dtypes(include="number").columns.tolist()
    if numeric_cols:
        total = pd.to_numeric(df[numeric_cols[0]], errors="coerce").sum()
        doc.add_paragraph(f"金额列：{numeric_cols[0]}")
        doc.add_paragraph(f"总金额：{total:,.2f}")

    # 整体风险评级
    doc.add_heading("2. 整体风险评级", level=1)
    risk = insight.get("overall_risk", "低风险")
    risk_para = doc.add_paragraph()
    risk_run = risk_para.add_run(f"综合评估：{risk}")
    risk_run.bold = True
    risk_run.font.size = Pt(14)
    risk_colors = {"高风险": RGBColor(0xD3, 0x2F, 0x2F), "中风险": RGBColor(0xF9, 0xA8, 0x25), "低风险": RGBColor(0x19, 0x76, 0xD2)}
    risk_run.font.color.rgb = risk_colors.get(risk, RGBColor(0x00, 0x00, 0x00))

    # AI 洞察
    doc.add_heading("3. 关键审计洞察", level=1)
    for idx, item in enumerate(insight.get("insights", []), 1):
        doc.add_heading(f"3.{idx} {item.get('title', '未命名')}", level=2)
        doc.add_paragraph(f"风险等级：{item.get('risk_level', 'N/A')}")
        doc.add_paragraph(f"置信度：{item.get('confidence', 'N/A')}")
        doc.add_paragraph(f"分析内容：{item.get('content', '')}")
        doc.add_paragraph(f"审计建议：{item.get('recommendation', '')}")

    # 异常交易明细
    doc.add_heading("4. 异常交易明细", level=1)
    if anomalies.empty:
        doc.add_paragraph("未发现明显异常记录。")
    else:
        doc.add_paragraph(f"共发现 {len(anomalies)} 条异常记录，分布如下：")
        # 异常类型统计
        type_counts = anomalies["异常类型"].value_counts()
        for atype, cnt in type_counts.items():
            doc.add_paragraph(f"   • {atype}：{cnt} 条", style="List Bullet")

        # 风险等级统计
        if "风险等级" in anomalies.columns:
            risk_counts = anomalies["风险等级"].value_counts()
            doc.add_paragraph("风险等级分布：")
            for rl, cnt in risk_counts.items():
                doc.add_paragraph(f"   • {rl}：{cnt} 条", style="List Bullet")

        # 关键字段表格
        key_cols = []
        for c in ["日期", "date", "金额", "amount", "摘要", "描述", "科目", "account", "异常类型", "风险等级", "Z分数"]:
            matches = [col for col in anomalies.columns if c.lower() in col.lower()]
            key_cols.extend(matches)
        key_cols = list(dict.fromkeys(key_cols))[:8]  # 去重并限制列数

        if key_cols:
            display_df = anomalies[key_cols].head(100).copy()
            table = doc.add_table(rows=1, cols=len(display_df.columns))
            table.style = "Light Shading Accent 1"
            hdr = table.rows[0].cells
            for i, col in enumerate(display_df.columns):
                hdr[i].text = str(col)
            for _, row in display_df.iterrows():
                cells = table.add_row().cells
                for i, col in enumerate(display_df.columns):
                    val = str(row[col])
                    if len(val) > 80:
                        val = val[:77] + "..."
                    cells[i].text = val
            doc.add_paragraph(f"注：以上为前 100 条异常记录摘要，完整数据请参见导出文件。")

    # 趋势分析
    if trend.get("has_data"):
        doc.add_heading("5. 趋势分析摘要", level=1)
        doc.add_paragraph(f"分析月份数：{trend['evidence'].get('months_analyzed', 'N/A')}")
        if trend.get("spikes"):
            doc.add_paragraph(f"检测到 {len(trend['spikes'])} 个趋势突变点，请重点关注以下月份：")
            for sp in trend["spikes"][:5]:
                doc.add_paragraph(f"   • {sp['month']}：金额 {sp['amount']:,.2f}，环比变化 {sp.get('change_pct', 0):.1%}", style="List Bullet")
        else:
            doc.add_paragraph("未发现显著趋势突变。")

    # 证据链附录
    if evidence_chain:
        doc.add_heading("附录：证据链记录", level=1)
        for idx, ev in enumerate(evidence_chain, 1):
            doc.add_heading(f"A.{idx} {ev.get('type', '未命名')}", level=2)
            doc.add_paragraph(f"样本量：{ev.get('sample_size', 'N/A')} 条")
            doc.add_paragraph(f"命中数：{ev.get('count', 'N/A')} 条")
            doc.add_paragraph(f"公式：{ev.get('formula', 'N/A')}")
            doc.add_paragraph(f"假设：{ev.get('assumption', 'N/A')}")

    # 免责声明
    doc.add_paragraph()
    p = doc.add_paragraph(EY_DISCLAIMER)
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_html_report(
    df: pd.DataFrame,
    anomalies: pd.DataFrame,
    trend: Dict[str, Any],
    insight: Dict[str, Any],
    evidence_chain: List[Dict[str, Any]],
    project_info: Dict[str, str],
    clean_log: Optional[Dict[str, Any]] = None,
) -> str:
    """生成 HTML 客户交付报告"""
    from memo_generator import generate_html_memo

    df_summary = {
        "total_rows": len(df),
        "clean_rows": len(df),
        "amount_col": df.select_dtypes(include="number").columns[0] if len(df.select_dtypes(include="number").columns) > 0 else None,
        "date_col": None,
        "total_amount": 0,
    }
    if df_summary["amount_col"]:
        df_summary["total_amount"] = pd.to_numeric(df[df_summary["amount_col"]], errors="coerce").sum()

    for col in df.columns:
        if any(k in col.lower() for k in ["日期", "date", "时间"]):
            df_summary["date_col"] = col
            break

    html = generate_html_memo(
        project_name=project_info.get("name", "未命名项目"),
        period=project_info.get("period", "未指定期间"),
        df_summary=df_summary,
        insight=insight,
        evidence_chain=evidence_chain,
        clean_log=clean_log or {},
    )
    return html
