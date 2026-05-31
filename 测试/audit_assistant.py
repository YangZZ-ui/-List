import streamlit as st
import pandas as pd
import plotly.express as px
import re
from datetime import datetime
from io import BytesIO
import io

# Word 报告依赖（安装提示在外层处理）
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ---------- 页面配置 ----------
st.set_page_config(page_title="智能审计洞察助手", layout="wide")
st.title("📊 智能审计洞察助手")
st.markdown("自然语言查询 · 异常检测 · 可视化看板 · 一键生成客户报告 · 导出全部数据")

# ---------- 会话状态 ----------
if "df" not in st.session_state:
    st.session_state.df = None
if "file_name" not in st.session_state:
    st.session_state.file_name = None
if "anomalies" not in st.session_state:
    st.session_state.anomalies = pd.DataFrame()
if "query_scope" not in st.session_state:
    st.session_state.query_scope = "全部数据"
if "nl_query" not in st.session_state:
    st.session_state.nl_query = ""

# ---------- 数据加载 ----------
def load_data(uploaded_file):
    try:
        if uploaded_file.name.endswith(".csv"):
            try:
                df = pd.read_csv(uploaded_file, encoding='utf-8-sig')
            except UnicodeDecodeError:
                uploaded_file.seek(0)
                df = pd.read_csv(uploaded_file, encoding='gbk')
        else:
            df = pd.read_excel(uploaded_file, engine="openpyxl")
        df.columns = df.columns.str.strip()
        return df
    except Exception as e:
        st.error(f"文件读取失败：{e}")
        return None

with st.sidebar:
    st.header("1. 数据上传")
    uploaded_file = st.file_uploader("上传试算平衡表/序时账/费用明细（Excel/CSV）", type=["xlsx", "csv"])
    if uploaded_file:
        st.session_state.df = load_data(uploaded_file)
        st.session_state.file_name = uploaded_file.name
        st.success(f"已加载 {len(st.session_state.df)} 行数据")
    
    st.header("2. 数据导出")
    if st.session_state.df is not None:
        # 导出全部原始数据
        csv = st.session_state.df.to_csv(index=False).encode('utf-8-sig')
        st.download_button(
            label="📥 导出全部原始数据 (CSV)",
            data=csv,
            file_name=f"{st.session_state.file_name}_全部数据.csv",
            mime="text/csv",
        )
        # 可选：导出Excel
        buffer = BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            st.session_state.df.to_excel(writer, index=False, sheet_name='原始数据')
            if not st.session_state.anomalies.empty:
                st.session_state.anomalies.to_excel(writer, index=False, sheet_name='异常数据')
        buffer.seek(0)
        st.download_button(
            label="📊 导出全部数据 (Excel，包含双sheet)",
            data=buffer,
            file_name=f"{st.session_state.file_name}_全部数据.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

# ---------- 辅助函数 ----------
def detect_amount_col(df):
    for col in df.columns:
        col_lower = col.lower()
        if any(k in col_lower for k in ["金额", "amount", "借方", "贷方", "balance", "amt", "本地货币"]):
            return col
    for col in df.columns:
        if pd.api.types.is_numeric_dtype(df[col]) and col not in ["年", "月", "日"]:
            return col
    return None

def detect_date_col(df):
    for col in df.columns:
        if any(k in col.lower() for k in ["日期", "date", "记账日期", "凭证日期", "posting"]):
            return col
    return None

def detect_text_col(df):
    for col in df.columns:
        if any(k in col.lower() for k in ["摘要", "描述", "text", "header", "说明"]):
            return col
    return None

def detect_account_col(df):
    for col in df.columns:
        col_lower = col.lower()
        if any(k in col_lower for k in ["科目", "account", "gl", "代码", "assignment"]):
            return col
    return None

# ---------- 增强的自然语言查询（支持选择数据范围）----------
def parse_nl_query(df, anomalies_df, query_text, scope="全部数据"):
    """
    scope: "全部数据" 或 "异常数据"
    """
    # 根据范围选择数据源
    if scope == "异常数据" and not anomalies_df.empty:
        data_df = anomalies_df.copy()
        scope_desc = "异常数据"
    else:
        data_df = df.copy()
        scope_desc = "全部数据"
        if scope == "异常数据" and anomalies_df.empty:
            st.warning("当前没有异常数据，已切换为查询全部数据")
    
    query_original = query_text.strip()
    query_lower = query_original.lower()
    amount_col = detect_amount_col(data_df)
    date_col = detect_date_col(data_df)
    account_col = detect_account_col(data_df)
    text_col = detect_text_col(data_df)

    # 1. 最高/最低金额（支持自定义笔数）
    if "最高" in query_lower or "最大" in query_lower:
        if amount_col:
            # 提取数字，如“最高的20笔” -> 20，默认15
            n = 15
            match_n = re.search(r"(\d+)\s*笔", query_lower)
            if match_n:
                n = int(match_n.group(1))
            result = data_df.nlargest(n, amount_col)
            return result, f"在{scope_desc}中，金额最大的 {len(result)} 笔交易"
    if "最低" in query_lower or "最小" in query_lower:
        if amount_col:
            n = 15
            match_n = re.search(r"(\d+)\s*笔", query_lower)
            if match_n:
                n = int(match_n.group(1))
            result = data_df.nsmallest(n, amount_col)
            return result, f"在{scope_desc}中，金额最小的 {len(result)} 笔交易"

    # 2. 节假日凭证
    if "节假日" in query_lower and date_col:
        if pd.api.types.is_datetime64_any_dtype(data_df[date_col]):
            holidays = ["12-31", "01-01", "04-04", "04-05", "05-01", "05-02", "05-03", "06-10", "06-11", "09-17", "09-18", "10-01", "10-02", "10-03", "10-04", "10-05", "10-06", "10-07"]
            data_df_temp = data_df.copy()
            data_df_temp["月日"] = data_df_temp[date_col].dt.strftime("%m-%d")
            result = data_df_temp[data_df_temp["月日"].isin(holidays)].drop(columns=["月日"])
            return result, f"在{scope_desc}中，节假日记账记录，共 {len(result)} 条"
    
    # 3. 处理“列名 等于 值” 或 “列名 = 值” 或 “列名 == 值”
    eq_pattern = r'(.+?)\s*(?:等于|==|=)\s*(.+)'
    match_eq = re.search(eq_pattern, query_original)
    if match_eq:
        col_part = match_eq.group(1).strip()
        value_part = match_eq.group(2).strip()
        if (value_part.startswith('"') or value_part.startswith("'")) and (value_part.endswith('"') or value_part.endswith("'")):
            value_part = value_part[1:-1]
        
        matched_col = None
        if col_part in data_df.columns:
            matched_col = col_part
        else:
            for col in data_df.columns:
                if col_part.lower() in col.lower():
                    matched_col = col
                    break
            if not matched_col and "account" in col_part.lower():
                matched_col = detect_account_col(data_df)
        
        if matched_col:
            try:
                val_num = float(value_part)
                if pd.api.types.is_numeric_dtype(data_df[matched_col]):
                    result = data_df[data_df[matched_col] == val_num]
                else:
                    result = data_df[data_df[matched_col].astype(str).str.strip() == value_part]
            except ValueError:
                result = data_df[data_df[matched_col].astype(str).str.strip() == value_part]
            if not result.empty:
                return result, f"在{scope_desc}中，筛选 {matched_col} = {value_part}，共 {len(result)} 条"
            else:
                return pd.DataFrame(), f"在{scope_desc}中未找到 {matched_col} = {value_part} 的记录"
    
    # 4. 处理“列名 包含 关键词”
    contain_pattern = r'(.+?)\s*包含\s*(.+)'
    match_cont = re.search(contain_pattern, query_original)
    if match_cont:
        col_part = match_cont.group(1).strip()
        keyword = match_cont.group(2).strip()
        matched_col = None
        if col_part in data_df.columns:
            matched_col = col_part
        else:
            for col in data_df.columns:
                if col_part.lower() in col.lower():
                    matched_col = col
                    break
        if matched_col:
            result = data_df[data_df[matched_col].astype(str).str.contains(keyword, na=False, case=False)]
            if not result.empty:
                return result, f"在{scope_desc}中，筛选 {matched_col} 包含 '{keyword}'，共 {len(result)} 条"
            else:
                return pd.DataFrame(), f"在{scope_desc}中未找到 {matched_col} 包含 '{keyword}' 的记录"
    
    # 5. 大于/小于
    match_gt = re.search(r"大于\s*([0-9,.]+)", query_lower)
    if match_gt and amount_col:
        val = float(match_gt.group(1).replace(',', ''))
        result = data_df[data_df[amount_col] > val]
        return result, f"在{scope_desc}中，{amount_col} > {val:,.2f}，共 {len(result)} 条"
    match_lt = re.search(r"小于\s*([0-9,.]+)", query_lower)
    if match_lt and amount_col:
        val = float(match_lt.group(1).replace(',', ''))
        result = data_df[data_df[amount_col] < val]
        return result, f"在{scope_desc}中，{amount_col} < {val:,.2f}，共 {len(result)} 条"
    
    # 6. 关键词在摘要/科目中
    if text_col:
        keywords = ["收入", "费用", "销售", "管理", "工资", "租金", "transfer", "salary"]
        matched_kw = [kw for kw in keywords if kw in query_lower]
        if matched_kw:
            result = data_df[data_df[text_col].str.contains("|".join(matched_kw), na=False, case=False)]
            if not result.empty:
                return result, f"在{scope_desc}中，筛选 {text_col} 包含 {', '.join(matched_kw)}，共 {len(result)} 条"
            else:
                return pd.DataFrame(), f"在{scope_desc}中未找到 {text_col} 包含相关关键词的记录"
    
    # 7. 默认：返回前100行，并提示
    return data_df.head(100), f"在{scope_desc}中未识别到明确条件，显示前100行（共{len(data_df)}行）"

# ---------- 异常检测 ----------
def detect_anomalies(df, amount_col, date_col=None):
    anomalies = pd.DataFrame()
    if amount_col is None:
        return anomalies
    df_work = df.copy()
    df_work[amount_col] = pd.to_numeric(df_work[amount_col], errors='coerce')
    df_work = df_work.dropna(subset=[amount_col])
    
    # 大额整数金额（仅标记超过中位数的整数金额，减少小额误报）
    if len(df_work) > 0:
        median_amt = df_work[amount_col].abs().median()
        int_mask = df_work[amount_col].astype(str).str.match(r'^-?\d+\.?0*$', na=False)
        int_mask &= df_work[amount_col].abs() >= median_amt
        int_amt = df_work[int_mask]
        if len(int_amt) > 0:
            int_amt["异常类型"] = "大额整数金额"
            anomalies = pd.concat([anomalies, int_amt], ignore_index=True)
    
    # 极端值
    if len(df_work) > 1:
        mean_val = df_work[amount_col].mean()
        std_val = df_work[amount_col].std()
        if std_val > 0:
            extreme = df_work[abs(df_work[amount_col] - mean_val) > 3 * std_val]
            if len(extreme) > 0:
                extreme["异常类型"] = "统计极端值"
                anomalies = pd.concat([anomalies, extreme], ignore_index=True)
    
    # 节假日
    if date_col and pd.api.types.is_datetime64_any_dtype(df_work[date_col]):
        holidays = ["12-31", "01-01", "04-04", "04-05", "05-01", "05-02", "05-03", "06-10", "06-11", "09-17", "09-18", "10-01", "10-02", "10-03", "10-04", "10-05", "10-06", "10-07"]
        df_work["月日"] = df_work[date_col].dt.strftime("%m-%d")
        holiday_rows = df_work[df_work["月日"].isin(holidays)]
        if len(holiday_rows) > 0:
            holiday_rows["异常类型"] = "节假日凭证"
            anomalies = pd.concat([anomalies, holiday_rows], ignore_index=True)
        df_work.drop(columns=["月日"], inplace=True)
    
    # 按所有原始列去重，避免仅按金额去重导致丢失不同交易记录
    anomalies = anomalies.drop_duplicates(keep='first')
    return anomalies

# ---------- Word报告生成（精简版，基于异常数据）----------
def generate_word_report(df, anomalies, amount_col):
    if not HAS_DOCX:
        raise ImportError("python-docx 未安装，请运行：pip install python-docx")

    doc = Document()
    title = doc.add_heading('审计数据洞察报告', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'数据源：{st.session_state.get("file_name", "未命名文件")}')
    
    doc.add_heading('1. 数据概览', level=2)
    doc.add_paragraph(f'原始数据总行数：{len(df):,}')
    if amount_col:
        try:
            total_amount = pd.to_numeric(df[amount_col], errors='coerce').sum()
            doc.add_paragraph(f'金额列：{amount_col}')
            doc.add_paragraph(f'总金额：{total_amount:,.2f}')
        except Exception:
            doc.add_paragraph(f'金额列：{amount_col}（计算总和时出错）')
    
    doc.add_heading('2. 异常交易摘要', level=2)
    if anomalies.empty:
        doc.add_paragraph("✅ 未发现明显异常。")
    else:
        anomaly_counts = anomalies['异常类型'].value_counts()
        doc.add_paragraph(f'⚠️ 共发现 {len(anomalies)} 条异常记录，分布如下：')
        for atype, cnt in anomaly_counts.items():
            doc.add_paragraph(f'   • {atype}：{cnt} 条', style='List Bullet')
        
        doc.add_heading('3. 异常交易明细（关键字段）', level=2)
        # 精简列
        key_cols = []
        date_col = detect_date_col(anomalies)
        if date_col:
            key_cols.append(date_col)
        if amount_col and amount_col in anomalies.columns:
            key_cols.append(amount_col)
        if '异常类型' in anomalies.columns:
            key_cols.append('异常类型')
        text_col = detect_text_col(anomalies)
        if text_col:
            key_cols.append(text_col)
        acc_col = detect_account_col(anomalies)
        if acc_col:
            key_cols.append(acc_col)
        for col in anomalies.columns:
            if any(k in col.lower() for k in ['document', '凭证', '编号', 'number']):
                if col not in key_cols:
                    key_cols.append(col)
                    break
        if len(key_cols) < 3:
            for col in anomalies.columns:
                if col not in key_cols and col != '异常类型':
                    key_cols.append(col)
                    if len(key_cols) >= 5:
                        break
        display_df = anomalies[key_cols].copy()
        if amount_col and amount_col in display_df.columns:
            display_df[amount_col] = display_df[amount_col].apply(lambda x: f"{x:,.2f}" if pd.notnull(x) else "")
        if date_col and date_col in display_df.columns:
            display_df[date_col] = pd.to_datetime(display_df[date_col], errors='coerce').dt.strftime("%Y-%m-%d")
        
        table = doc.add_table(rows=1, cols=len(display_df.columns))
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(display_df.columns):
            hdr_cells[i].text = str(col)
        for _, row in display_df.iterrows():
            row_cells = table.add_row().cells
            for i, col in enumerate(display_df.columns):
                val = str(row[col])
                if len(val) > 80:
                    val = val[:77] + "..."
                row_cells[i].text = val
        doc.add_paragraph()
        doc.add_paragraph(f'注：以上为异常交易的关键信息摘要。完整异常记录共 {len(anomalies)} 行。')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------- 主界面 ----------
if st.session_state.df is not None:
    df = st.session_state.df
    amount_col = detect_amount_col(df)
    date_col = detect_date_col(df)
    
    # 移除原始数据自动预览，改为一个可选的按钮
    if st.button("📂 预览原始数据（前100行）"):
        st.dataframe(df.head(100))
    
    st.subheader("🔍 自然语言数据查询")
    # 查询范围选择
    col_scope, _ = st.columns([1, 3])
    with col_scope:
        query_scope = st.radio(
            "查询范围",
            ["全部数据", "异常数据"],
            index=0 if st.session_state.query_scope == "全部数据" else 1,
            horizontal=True,
            key="query_scope_radio"
        )
        st.session_state.query_scope = query_scope
    
    col1, col2 = st.columns([3,1])
    with col1:
        query = st.text_input(
            "输入你的问题，例如：金额大于10000的记录 / 展示最高的15笔金额 / 查找销售费用 / Account等于91100",
            value=st.session_state.nl_query,
            key="nl_query_input"
        )
        # 同步 session_state 与当前输入
        st.session_state.nl_query = query
    with col2:
        st.write("快速提问：")
        if st.button("💰 金额最高15笔"):
            st.session_state.nl_query = "金额最高15笔"
            st.rerun()
        if st.button("📅 节假日记账"):
            st.session_state.nl_query = "节假日凭证"
            st.rerun()
    
    if query:
        # 根据选择的范围进行查询
        if query_scope == "异常数据" and st.session_state.anomalies.empty:
            st.warning("当前没有异常数据，请先点击「开始检测异常」")
        else:
            result, explain = parse_nl_query(df, st.session_state.anomalies, query, scope=query_scope)
            st.info(f"查询解析：{explain}")
            if result.empty:
                st.warning("未查询到符合条件的记录，请检查条件或数据。")
            else:
                st.dataframe(result)
                # 如果结果不为空且有金额列，显示金额分布图
                if amount_col and len(result) > 0:
                    fig = px.bar(result, x=result.index, y=amount_col, title=f"查询结果金额分布 (共{len(result)}条)")
                    st.plotly_chart(fig, use_container_width=True)
    
    st.subheader("⚠️ 异常交易检测")
    if st.button("开始检测异常"):
        with st.spinner("正在分析数据..."):
            anomalies = detect_anomalies(df, amount_col, date_col)
            if not anomalies.empty:
                st.warning(f"发现 {len(anomalies)} 条异常记录")
                st.dataframe(anomalies)
                st.session_state.anomalies = anomalies
            else:
                st.success("未发现明显异常")
                st.session_state.anomalies = pd.DataFrame()
    
    st.subheader("📈 可视化看板")
    tab1, tab2, tab3 = st.tabs(["金额分布", "异常饼图", "时间趋势"])
    with tab1:
        if amount_col:
            fig_hist = px.histogram(df, x=amount_col, nbins=30, title="金额分布直方图")
            st.plotly_chart(fig_hist, use_container_width=True)
    with tab2:
        if not st.session_state.anomalies.empty:
            anomaly_counts = st.session_state.anomalies['异常类型'].value_counts()
            fig_pie = px.pie(values=anomaly_counts.values, names=anomaly_counts.index, title="异常类型构成")
            st.plotly_chart(fig_pie, use_container_width=True)
        else:
            st.info("点击「开始检测异常」生成异常数据")
    with tab3:
        if date_col and pd.api.types.is_datetime64_any_dtype(df[date_col]):
            df_month = df.set_index(date_col).resample('M')[amount_col].sum().reset_index()
            fig_line = px.line(df_month, x=date_col, y=amount_col, title="月度金额趋势")
            st.plotly_chart(fig_line, use_container_width=True)
        else:
            st.info("缺少日期列或日期格式不正确")
    
    st.subheader("📄 生成客户沟通草稿（基于异常数据）")
    if st.button("一键生成Word报告"):
        if not HAS_DOCX:
            st.error("未安装 python-docx，请运行：pip install python-docx")
        else:
            anomalies = st.session_state.get('anomalies', pd.DataFrame())
            try:
                buffer = generate_word_report(df, anomalies, amount_col)
                st.download_button(
                    label="📥 下载Word报告",
                    data=buffer,
                    file_name="audit_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"报告生成失败：{e}")
else:
    st.info("请从左侧上传数据文件开始")